#表格字体颜色

from tkinter import *
from tkinter.filedialog import askdirectory
import tkinter.messagebox
import tkinter as tk


from tkinter import filedialog

from tkinter.filedialog import askdirectory

#C://Users//Administrator//Downloads//表格自动着色测试.docx
from  docx import  Document
from docx.shared import RGBColor


def get_info(document):
    tables = document.tables
    paras =  document.paragraphs
    str1 = ''
    for para in paras:
        str1 = str1 + para.text

    role = r'///(.*?)///'
    request_li = re.findall(role, str1)
    print(request_li)

    request_list = []
    for request in request_li:
        request_dir = {}
        try:
            request_each = request.split(';')
            print('a', request_each)

            if len(request_each) != 6:
                print('第 %s个表格的需求有缺失必要的参数 ' % (request_li.index(request) + 1))
            for request_info in request_each[0:5]:
                request_info = request_info.split(':')
                # request_info_ =  request_info.split(':')
                request_dir[request_info[0]] = eval(request_info[1])
            request_list.append(request_dir)
        except:
           print('第 %s个表格的需求格式有误'%(request_li.index(request)+1))

    if  not len(request_list) == len(tables):
        print('请核对是否有表格未添加需求')

    return request_list,tables

#B;lue: 0 255 0 -->red : 255 0 0
def get_colors():
    step = 25
    color_list =[]
    r, g, b = 0, 255, 0
    for i in range(10):
        r,g,b = r+step,g-step,b
        color_list.append((r,g,b))
    print('color',color_list)
    return color_list



def get_number_step(start,end):

    step = (end - start)/10
    num_list  =[]
    lift ,right = start,start+step
    num_list.append((lift,right))

    for i in range(9):

        lift,right  = right , right+step,
        if i ==8:
            right = end +1
        num_list.append((lift,right))
    print(num_list)
    return num_list

def makeup_for_table(request_list,tables,color_list):
    table_num = 0
    try:
       for table in tables:
            request_dir_table = request_list[table_num]
            print('dir',request_dir_table)
            num_range = request_dir_table['Color']
            Start_row = request_dir_table['Start_row']
            Start_col = request_dir_table['Start_col']
            end_row = request_dir_table['end_row']
            end_col = request_dir_table['end_col']
            # deal with the number step
            start, end = num_range[0], num_range[1]
            num_step_list = get_number_step(start, end)
            print('num_step_list',num_step_list)
            print(num_step_list)
            for r in range(Start_row - 1, len(table.rows) + end_row + 1):
                for c in range(Start_col - 1, len(table.columns) + end_col + 1):
                    text = table.cell(r, c).text
                    print(text)
                    num = int(text)
                    for i in range(10):
                        if num_step_list[i][0] <= num < num_step_list[i][1]:
                            color_for_cell = color_list[i]

                    table.cell(r, c).paragraphs[0].clear()
                    my_paragraph = table.cell(r, c).paragraphs[0]
                    run2 = my_paragraph.add_run(text)
                    color = RGBColor(color_for_cell[0], color_for_cell[1], color_for_cell[2])
                    run2.font.color.rgb = color
            table_num +=1
    except:
        print("%s 表格着色失败"%table_num+1)



def choose_file():
    root = tkinter.Tk()
    root.withdraw()

    file_path = filedialog.askopenfilename(filetypes=[('doc', '*.docx')])
    return file_path

# 目录查询
def selectPath():
    path_ = askdirectory()
    path.set(path_)

if __name__ == '__main__':
    while True:

        input('按Enter选择要着色的文件')
        path = choose_file()
        print(path)
        try:
            document = Document(path)
        except:
            print("请确认要输入的文件路径")
            choose = input('输入 1 ： 继续，输入其他 ： 退出')
            try:
                choose = int(choose)
                if choose == 1:
                    continue
                else:
                    break
            except:
                break
        try:
            request_list ,tables = get_info(document)
            color_list = get_colors()
            makeup_for_table(request_list,tables,color_list)
        except:
            # print('程序运行失败。请确认文档格式是否符合标准')
            choose = input('输入 1 ： 继续，输入其他 ： 退出')
            try:
                choose = int(choose)
                if choose == 1:
                    continue
                else:
                    break
            except:
                break

        try:
            input('按Enter ：文件将自动保存为原文件路径下：文件名+着色版.docx文件')
            # import_path  = input('input the filename to save as : ')
            save_path = path.split('.docx')[0]+'着色版'+'.docx'
            print(save_path)
            document.save(save_path)
            print('文件保存成功')
        except:
            print('保存失败，请先关闭输出文件')
        choose = input('输入 1 ： 继续:输入其他 ： 退出')
        try:
            choose = int(choose)
            if choose == 1:
                continue
            else:
                break
        except:
            break














