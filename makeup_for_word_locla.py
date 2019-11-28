
from docx.oxml import parse_xml
import re
import pandas as pd
from docx import Document
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor



def get_info(document):
    tables = document.tables
    paras = document.paragraphs
    str1 = ''
    for para in paras:
        str1 = str1 + para.text

    role = r'///(.*?)///'
    request_li = re.findall(role, str1)
    str1= ''
    str2= ''
    str3= ''
    request_list = []
    for request in request_li:
        request_dir = {}
        try:
            request_each = request.split(';')
            if len(request_each) != 6:
                str1 = '第 %s个表格的需求有缺失必要的参数 ' % (request_li.index(request) + 1)
                str1 = str.encode('utf-8').decode('utf-8')
            for request_info in request_each[0:5]:
                request_info = request_info.split(':')
                request_dir[request_info[0]] = eval(request_info[1])
            request_list.append(request_dir)
        except:
            str2 = '第 %s个表格的需求格式有误' % (request_li.index(request) + 1)
            str2 = str2.encode('utf-8').decode('utf-8')

    if not len(request_list) == len(tables):
        str3  = '请核对是否有表格未添加需求'
    errmsg = str1+str2+str3

    return request_list, tables,errmsg


# B;lue: 0 255 0 -->red : 255 0 0
def get_colors():
    color_list = [
        "20B2AA",
        "90EE90",
        "00FFFF",
        "ADD8FF",
        "ADD8E6",
        "F0E68C",
        "FFFFE0",
        "FFD700",
        "FFA500",
        "FF69B4",
    ]
    return color_list


def get_number_step(start, end):
    step = (end - start) / 10
    num_list = []
    lift, right = start, start + step
    num_list.append((lift, right))

    for i in range(9):

        lift, right = right, right + step,
        if i == 8:
            right = end + 1
        num_list.append((lift, right))
    return num_list


def makeup_for_table(request_list, tables, color_list):
    table_num = 0
    try:
        for table in tables:
            request_dir_table = request_list[table_num]
            num_range = request_dir_table['Color']
            Start_row = request_dir_table['Start_row']
            Start_col = request_dir_table['Start_col']
            end_row = request_dir_table['end_row']
            end_col = request_dir_table['end_col']

            #將數據分為10組
            start, end = num_range[0], num_range[1]
            num_step_list = get_number_step(start, end)
            #背景著色
            for r in range(Start_row - 1, len(table.rows) + end_row + 1):
                for c in range(Start_col - 1, len(table.columns) + end_col + 1):
                    text = table.cell(r, c).text
                    num = int(text)
                    for i in range(10):
                        if num_step_list[i][0] <= num < num_step_list[i][1]:
                            color_for_cell = color_list[i]
                    shading_elm_1 = parse_xml(
                        r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=color_for_cell))

                    table.cell(r, c)._tc.get_or_add_tcPr().append(shading_elm_1)
            table_num += 1
        msg = ''
        return  msg
    except Exception as e :
        msg  = '表格轉換失敗： ' +e
        return msg




import sys

if __name__ == '__main__':
    message = input()

    message = message.replace('\n', '')
    # message = message.encode('gbk').decode('gbk')
    path = message.split(';')

    file_path = path[0]

    document = Document(file_path)

    save_path = path[1] + path[2] + '着色版' + '.docx'

    request_list, tables ,errmsg= get_info(document)
    color_list = get_colors()
    msg =  makeup_for_table(request_list, tables, color_list)
    document.save(save_path)
    #
    # with open(save_path,'w') as f:
    #     print('hello', file=f, flush=True)
    print('info:',errmsg+msg,'file_path:', save_path)

